from __future__ import annotations

import csv
import json
import re
import shutil
from pathlib import Path
from typing import Iterable

import openpyxl
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "06. Ataskaitos knygutė 2026"
PROJECT = ROOT / "ataskaita-2026"


DOC_MAP = {
    "[1] Sveikinimo žodžiai/[1.].docx": "src/sveikinimai.md",
    "[2] Veiklos principai ir valdymas/[2.1] VU SA Strategija.docx": "src/strategija.md",
    "[2] Veiklos principai ir valdymas/[2.2] VU SA Parlamentas.docx": "src/vu-sa/parlamentas.md",
    "[2] Veiklos principai ir valdymas/[2.3] VU SA Taryba.docx": "src/vu-sa/taryba.md",
    "[2] Veiklos principai ir valdymas/[2.4] VU SA Revizija.docx": "src/vu-sa/revizijos-komisija.md",
    "[2] Veiklos principai ir valdymas/[2.5] VU SA ISF.docx": "src/vu-sa/isf.md",
    "[2] Veiklos principai ir valdymas/[2.6] VU SA DAG.docx": "src/vu-sa/dag.md",
    "[3] Metų veiklos planas/[3.1] Strateginiai tikslai.docx": "src/mvp/strateginiai.md",
    "[3] Metų veiklos planas/[3.2] VU SA P tikslai.docx": "src/mvp/tikslai-padaliniuose.md",
    "[4] Kokybiškos studijos ir joms pritaikyta aplinka/[4.1] Nacionalinis ir miesto savivaldoje atstovavimas.docx": "src/kokybiskos-studijos/nacionalinis.md",
    "[4] Kokybiškos studijos ir joms pritaikyta aplinka/[4.2] Atstovavimas VU.docx": "src/kokybiskos-studijos/vu.md",
    "[4] Kokybiškos studijos ir joms pritaikyta aplinka/[4.3] Bendradarbiavimas.docx": "src/kokybiskos-studijos/bendradarbiavimas.md",
    "[4] Kokybiškos studijos ir joms pritaikyta aplinka/[4.4] Tarptautinis bendradarbiavimas ir konferencijos.docx": "src/kokybiskos-studijos/tarptautinis.md",
    "[4] Kokybiškos studijos ir joms pritaikyta aplinka/[4.5] Raštai.docx": "src/kokybiskos-studijos/rastai.md",
    "[5] Stipri organizacija/[5.1] Tyrimai.docx": "src/stipri-organizacija/tyrimai.md",
    "[5] Stipri organizacija/[5.2] Projektai.docx": "src/stipri-organizacija/projektai.md",
    "[5] Stipri organizacija/[5.3] VU SA procesų atnaujinimas.docx": "src/stipri-organizacija/procesu-atnaujinimas.md",
    "[5] Stipri organizacija/[5.4] INSTITUCINIS STIPRINIMAS.docx": "src/stipri-organizacija/isf.md",
    "[6] Darni universitetinė bendruomenė/[6.1] Bendruomenę buriantys renginiai.docx": "src/darni-bendruomene/renginiai.md",
    "[6] Darni universitetinė bendruomenė/[6.2] Integracija.docx": "src/darni-bendruomene/integracija.md",
    "[6] Darni universitetinė bendruomenė/[6.3] Studentų (-čių) iniciatyvos.docx": "src/darni-bendruomene/iniciatyvos.md",
}

PUBLIC_YEAR = "2025-2026"


def iter_block_items(parent):
    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def clean_text(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def para_text(paragraph: Paragraph) -> str:
    return clean_text(paragraph.text)


def is_heading(text: str, index: int, style_name: str) -> int | None:
    if not text:
        return None
    if index == 0:
        return 1
    style = style_name.lower()
    if "heading 1" in style or "antraštė 1" in style:
        return 1
    if "heading 2" in style or "antraštė 2" in style:
        return 2
    if "heading 3" in style or "antraštė 3" in style:
        return 3
    if len(text) <= 90 and text == text.upper() and re.search(r"[A-ZĄČĘĖĮŠŲŪŽ]", text):
        return 2
    if len(text) <= 110 and not text.endswith((".", ":", ";", ",")) and len(text.split()) <= 10:
        return 2
    return None


def paragraph_to_md(paragraph: Paragraph, index: int) -> list[str]:
    text = para_text(paragraph)
    if not text:
        return []
    style_name = paragraph.style.name if paragraph.style is not None else ""
    level = is_heading(text, index, style_name)
    if level:
        return [f"{'#' * level} {text}", ""]
    if style_name.lower().startswith(("list", "sąraš")):
        return [f"- {text}", ""]
    return [text, ""]


def table_to_html(table: Table) -> list[str]:
    lines = ['<table class="report-table">']
    for row in table.rows:
        cells = [clean_text(cell.text).replace("\n", "<br>") for cell in row.cells]
        tag = "th" if row is table.rows[0] else "td"
        lines.append("  <tr>" + "".join(f"<{tag}>{cell}</{tag}>" for cell in cells) + "</tr>")
    lines.append("</table>")
    lines.append("")
    return lines


def docx_to_markdown(src: Path, title: str | None = None, language: str = "lt") -> str:
    doc = Document(str(src))
    lines: list[str] = []
    paragraph_index = 0
    include = language == "lt"
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = para_text(block)
            if text.strip().upper() == "[EN]":
                if language == "lt":
                    break
                include = True
                continue
            if not include:
                continue
            converted = paragraph_to_md(block, paragraph_index)
            if converted:
                paragraph_index += 1
                lines.extend(converted)
        elif isinstance(block, Table):
            if not include:
                continue
            lines.extend(table_to_html(block))
    while lines and not lines[-1]:
        lines.pop()
    first_title = title or next((line[2:] for line in lines if line.startswith("# ")), src.stem)
    body = "\n".join(lines)
    return f"---\ntitle: {json.dumps(first_title, ensure_ascii=False)}\n---\n\n{body}\n"


def write_markdown_pages() -> None:
    for src_rel, dst_rel in DOC_MAP.items():
        src = SOURCE / src_rel
        dst = PROJECT / dst_rel
        dst.parent.mkdir(parents=True, exist_ok=True)
        dst.write_text(docx_to_markdown(src, language="lt"), encoding="utf-8", newline="\n")

        en_dst = PROJECT / "src" / "en" / Path(dst_rel).relative_to("src")
        en_dst.parent.mkdir(parents=True, exist_ok=True)
        en_dst.write_text(docx_to_markdown(src, language="en"), encoding="utf-8", newline="\n")

    padaliniai = SOURCE / "Padalinių ataskaitos" / "Nuorodos.docx"
    appendix = PROJECT / "src" / "padaliniu-ataskaitos.md"
    appendix.write_text(docx_to_markdown(padaliniai, "Padalinių ataskaitos", "lt"), encoding="utf-8", newline="\n")
    en_appendix = PROJECT / "src" / "en" / "padaliniu-ataskaitos.md"
    en_appendix.parent.mkdir(parents=True, exist_ok=True)
    en_appendix.write_text(docx_to_markdown(padaliniai, "Unit Reports", "en"), encoding="utf-8", newline="\n")


def stats_from_doc() -> dict[str, list[dict[str, object]]]:
    doc = Document(str(SOURCE / "[0] Pradinis puslapis" / "[0.2] VU SA skaičiais.docx"))
    lines = [para_text(p) for p in doc.paragraphs if para_text(p)]
    items: list[dict[str, object]] = []
    icons = ["Star", "Building", "Lightbulb", "Users"]
    for line in lines:
        if len(items) >= 4 or line == "VU SR IN NUMBERS":
            break
        if line.startswith("37 metai"):
            items.append({
                "value": 37,
                "label": "metai",
                "description": "Vilniaus universiteto studentų (-čių) atstovavimo",
                "icon": icons[len(items)],
            })
        elif line.startswith("15 VU SA padalinių"):
            items.append({
                "value": 15,
                "label": "VU SA padalinių",
                "description": "veikiančių 3 miestuose (Vilniuje, Kaune ir Šiauliuose)",
                "icon": icons[len(items)],
            })
        elif line.startswith("27 studentiškos iniciatyvos"):
            items.append({
                "value": 27,
                "label": "studentiškos iniciatyvos",
                "description": "kuriančios studentų gerovę",
                "icon": icons[len(items)],
            })
        elif line.startswith("1487"):
            items.append({
                "value": 1487,
                "label": "nariai",
                "description": "dėl bendro tikslo veikiantys nariai",
                "icon": icons[len(items)],
            })
    if not items:
        items = [
            {"value": 37, "label": "metai", "description": "Vilniaus universiteto studentų (-čių) atstovavimo", "icon": "Star"},
            {"value": 15, "label": "VU SA padalinių", "description": "veikiančių 3 miestuose", "icon": "Building"},
        ]
    en_items: list[dict[str, object]] = []
    en_started = False
    for line in lines:
        if line == "VU SR IN NUMBERS":
            en_started = True
            continue
        if not en_started or len(en_items) >= 4:
            continue
        if line.startswith("37 years"):
            en_items.append({
                "value": 37,
                "label": "years",
                "description": "of representing Vilnius University students",
                "icon": icons[len(en_items)],
            })
        elif line.startswith("15 VU SR units"):
            en_items.append({
                "value": 15,
                "label": "VU SR units",
                "description": "operating in 3 cities (Vilnius, Kaunas, and Šiauliai)",
                "icon": icons[len(en_items)],
            })
        elif line.startswith("27 student initiatives"):
            en_items.append({
                "value": 27,
                "label": "student initiatives",
                "description": "creating student well-being",
                "icon": icons[len(en_items)],
            })
        elif line.startswith("1487"):
            en_items.append({
                "value": 1487,
                "label": "members",
                "description": "working toward a common goal",
                "icon": icons[len(en_items)],
            })
    return {"lt": items, "en": en_items or items}


def timeline_from_doc() -> dict[str, list[dict[str, str]]]:
    doc = Document(str(SOURCE / "[0] Pradinis puslapis" / "[0.1] Įsimintiniausi metų įvykiai.docx"))
    paras = [para_text(p) for p in doc.paragraphs if para_text(p)]
    def collect(language: str) -> list[dict[str, str]]:
        skip = {"Įsimintiniausi metų įvykiai", "Highlights of the year"}
        events: list[dict[str, str]] = []
        current: dict[str, str] | None = None
        include = language == "lt"
        for text in paras:
            if text.strip().upper() == "[EN]":
                if language == "lt":
                    break
                include = True
                continue
            if not include or text in skip:
                continue
            if len(text) <= 80 and not text.endswith((".", "!", "?")):
                if current and current["description"]:
                    events.append(current)
                current = {"title": text, "description": ""}
            elif current:
                current["description"] = clean_text((current["description"] + " " + text).strip())
        if current and current["description"]:
            events.append(current)
        return events[:12]

    return {"lt": collect("lt"), "en": collect("en")}


def write_json(path: Path, data: object) -> None:
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8", newline="\n")


def copy_gallery_and_data() -> None:
    src_dir = SOURCE / "[0] Pradinis puslapis" / "[0.3] Nuotraukos iš šių metų įvykių"
    dst_dir = PROJECT / "src" / "public" / "img" / "gallery" / PUBLIC_YEAR
    dst_dir.mkdir(parents=True, exist_ok=True)
    entries = []
    for index, src in enumerate(sorted(src_dir.iterdir())):
        if not src.is_file() or src.suffix.lower() not in {".jpg", ".jpeg", ".png"}:
            continue
        dst = dst_dir / src.name
        shutil.copy2(src, dst)
        entries.append({
            "name": src.stem,
            "src": f"/img/gallery/{PUBLIC_YEAR}/{src.name}",
            **({"featured": True} if index in {0, 8} else {}),
        })
    write_json(PROJECT / "src" / "data" / "gallery.json", {"lt": entries, "en": entries})


def export_people_csv() -> None:
    workbook = openpyxl.load_workbook(
        SOURCE / "[7] Bendruomenė - nuotraukos" / "2025-2026 m. m. kadencija.xlsx",
        data_only=True,
    )
    outputs = {
        "Dariniai": PROJECT / "src" / "data" / "dariniai.csv",
        "Padaliniai": PROJECT / "src" / "data" / "padaliniai.csv",
    }
    for sheet_name, dst in outputs.items():
        ws = workbook[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if sheet_name == "Dariniai":
            rows = rows[1:]
        header = [clean_text(str(value)) if value is not None else "" for value in rows[0]][:6]
        data_rows = rows[1:]
        dst.parent.mkdir(parents=True, exist_ok=True)
        with dst.open("w", encoding="utf-8-sig", newline="") as handle:
            writer = csv.writer(handle)
            writer.writerow(header)
            for row in data_rows:
                values = ["" if value is None else clean_text(str(value)) for value in row[:6]]
                if any(values):
                    writer.writerow(values)


def update_index() -> None:
    path = PROJECT / "src" / "index.md"
    text = path.read_text(encoding="utf-8")
    replacements = {
        "VU SA Metinė ataskaita 2024-2025": "VU SA Metinė ataskaita 2025-2026",
        "už 2024-2025 metus": "už 2025-2026 metus",
        "/VU_SA_Ataskaita_2024_2025.pdf": "/VU_SA_Ataskaita_2025_2026.pdf",
        "VU SA 2024–2025 m. svarbiausios veiklos": "VU SA 2025–2026 m. svarbiausios veiklos",
        "VU SA 2024-2025 metų ataskaitoje": "VU SA 2025-2026 metų ataskaitoje",
        "VU SA 2024–2025 metų ataskaitoje": "VU SA 2025–2026 metų ataskaitoje",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    path.write_text(text, encoding="utf-8", newline="\n")

    en_dir = PROJECT / "src" / "en"
    en_dir.mkdir(parents=True, exist_ok=True)
    en_index = en_dir / "index.md"
    en_index.write_text("""---
layout: home

hero:
  name: "VU SR report"
  text: "for 2025-2026"
  tagline: "VU SR report for 2025-2026"
  image:
    light:
      src: /img/logos/vusa-en-b.png
      alt: VU SR
    dark:
      src: /img/logos/vusa-en-w.png
      alt: VU SR
  actions:
    - theme: brand
      text: Learn about this year!
      link: /en/mvp/strateginiai
    - theme: alt
      text: Download PDF Report
      link: /VU_SR_Report_2025_2026.pdf
---

<script setup lang="ts">
import { computed } from 'vue';
import stats from "../data/stats.json"
import galleryImages from "../data/gallery.json"
import timelineData from "../data/timeline.json"
import people from "../data/dariniai.csv"

import VPButton from "vitepress/dist/client/theme-default/components/VPButton.vue";
import PersonAvatar from "@/PersonAvatar.vue";
import EnhancedNumberStatistic from "@/EnhancedNumberStatistic.vue";
import ImageMosaic from "@/ImageMosaic.vue";
import InteractiveTimeline from "@/InteractiveTimeline.vue";

import { Building, Star, Lightbulb, Users } from 'lucide-vue-next';

const iconMap = {
  Star,
  Building,
  Lightbulb,
  Users,
};

const timelineEvents = timelineData.en;
</script>

<section class="lg:px-2 px-1.5 isolate">
  <div class="mx-auto relative">
    <figure class="grid grid-cols-1 md:grid-cols-2 gap-4 items-center rounded-xl p-6">
      <figcaption class="my-4 px-6 text-left order-2 md:order-1 flex flex-col">
        <p class="text-md font-medium italic mb-4 leading-6!">
          “We will enter the upcoming academic year with newly approved strategic projects that touch every member of our community.”
        </p>
        <p class="font-bold mb-4 leading-6!">
          I invite you to explore the work done throughout the year in the VU SR 2025-2026 report.
        </p>
        <PersonAvatar class="mt-4" :src="`/img/people/${people[0]['Nuotraukos pavadinimas']}`" size="small">
          <p style="margin: 0" class="font-bold leading-5!">{{ people[0]['Vardas Pavardė'] }}</p>
          <p style="margin: 0" class="opacity-80 text-sm">{{ people[0]['Pareigos'] }}</p>
        </PersonAvatar>
      </figcaption>
      <video playsinline autoplay controls muted loop class="mx-auto order-1 md:order-2 z-20 shadow-lg rounded-lg aspect-video">
        <source src="/video/kleja-en.webm" type="video/webm">
        Your browser does not support the video tag.
      </video>
    </figure>
  </div>
</section>

<section class="lg:px-2 px-1.5 isolate my-12">
  <div class="max-w-6xl mx-auto">
    <div class="text-center">
      <h2 style="border: 0; padding: 0" class="font-bold mb-2">VU SR is:</h2>
    </div>
    <div class="grid grid-cols-1 sm:grid-cols-2 lg:grid-cols-4 gap-6 mt-10">
      <template v-for="stat in stats.en" :key="stat.label">
        <EnhancedNumberStatistic
          :end-number="stat.value"
          :title="stat.label"
          :icon="iconMap[stat.icon]"
        >
          {{ stat.description }}
        </EnhancedNumberStatistic>
      </template>
    </div>
  </div>
</section>

<section class="lg:px-2 px-1.5 isolate my-8 py-4">
  <div class="max-w-6xl mx-auto">
    <InteractiveTimeline
      title="VU SR 2025-2026 highlights"
      :events="timelineEvents"
    />
  </div>
</section>

<section class="lg:px-2 px-1.5 isolate my-8 py-4">
  <div class="max-w-6xl mx-auto">
    <ImageMosaic
      title="VU SR moments"
      :images="galleryImages.en"
    />
  </div>
</section>

<section class="lg:px-2 px-1.5 isolate my-20">
  <div class="max-w-5xl mx-auto text-center p-10 bg-gradient-to-br from-amber-500/10 via-white to-amber-500/5 dark:from-amber-900/20 dark:via-gray-800 dark:to-amber-900/10 rounded-2xl shadow-lg backdrop-blur-sm">
    <h2 class="text-3xl font-bold mb-4">Join the VU SR community</h2>
    <p class="text-lg mb-8 px-12">Every VU student can join VU SR.</p>
    <div class="flex flex-wrap justify-center gap-4 mt-8">
      <VPButton href="https://vusa.lt/tapk-nariu" text="Become a member" />
      <VPButton href="/VU_SR_Report_2025_2026.pdf" text="Download PDF" theme="brand" />
      <VPButton href="https://vusa.lt/lt/kontaktai/centrinis-biuras" text="Contact" theme="alt" />
    </div>
  </div>
</section>
""", encoding="utf-8", newline="\n")


def write_english_support_pages() -> None:
    src = ROOT / "ataskaita-2025" / "src" / "en" / "bendruomene.md"
    dst = PROJECT / "src" / "en" / "bendruomene.md"
    if src.exists():
        dst.parent.mkdir(parents=True, exist_ok=True)
        text = src.read_text(encoding="utf-8").replace("2024-2025", "2025-2026").replace("2024–2025", "2025–2026")
        text = text.replace("\nMore on VU SR Board [here](/en/vu-sa/taryba.md).\n", "\n")
        text = text.replace("\nMore on VU SR Parliament [here](/en/vu-sa/parlamentas.md).\n", "\n")
        dst.write_text(text, encoding="utf-8", newline="\n")

    (PROJECT / "src" / "en" / "padeka.md").write_text("""---
outline: deep
---

# Acknowledgements

## Report Project Team

The 2025-2026 VU SR report project was coordinated by: **Samanta Valiušaitytė**

- Web layout: **Martynas Matijošius**
- Lithuanian text editors: **Austėja Simonaitytė, Mila Mašauskaitė**
- Visual content selection and preparation: **Eivinas Zableckas**

Thank you to the VU SR Central Office and everyone who contributed content for the report.
""", encoding="utf-8", newline="\n")


def remove_empty_english_pages() -> None:
    for path in (PROJECT / "src" / "en").rglob("*.md"):
        if path.name in {"index.md", "bendruomene.md", "padeka.md"}:
            continue
        lines = [
            line.strip()
            for line in path.read_text(encoding="utf-8").splitlines()
            if line.strip() and line.strip() != "---" and not line.strip().startswith("title:")
        ]
        if not lines:
            path.unlink()


def update_configs() -> None:
    replacements = {
        ".vitepress/config.mts": {
            "https://ataskaita2025.vusa.lt": "https://ataskaita2026.vusa.lt",
            "ataskaita 2025": "ataskaita 2026",
            "metinė ataskaita 2024-2025": "metinė ataskaita 2025-2026",
            "metinė ataskaita 2026-2026": "metinė ataskaita 2025-2026",
        },
        ".vitepress/config/lt.ts": {
            "Ataskaita 2025": "Ataskaita 2026",
            "2024-2025 m. kadenciją": "2025-2026 m. kadenciją",
        },
        ".vitepress/config/en.ts": {
            "Report 2025": "Report 2026",
            "2024-2025 term": "2025-2026 term",
            "/en/darni-bendruomene/darni-bendruomene": "/en/darni-bendruomene/renginiai",
            "Friends and the report project team": "Report project team",
        },
        "pdf-report/README.md": {
            "Annual Report 2025": "Annual Report 2026",
            "VU SA Annual Report 2025": "VU SA Annual Report 2026",
        },
    }
    for rel, reps in replacements.items():
        path = PROJECT / rel
        text = path.read_text(encoding="utf-8")
        for old, new in reps.items():
            text = text.replace(old, new)
        if rel == ".vitepress/config/en.ts":
            text = text.replace("            { text: 'Structure', link: '/en/vu-sa/parlamentas' },\n", "")
            text = re.sub(r"            \{\n                text: 'VU SR structure and principles of activity ⭐️',\n                items: \[\n(?:.*\n){6}                \]\n            \},\n", "", text)
            text = text.replace("                            { text: 'Cooperation', link: '/en/kokybiskos-studijos/bendradarbiavimas' },\n", "")
            text = text.replace("                            { text: 'Student initiatives', link: '/en/darni-bendruomene/iniciatyvos' },\n", "")
        path.write_text(text, encoding="utf-8", newline="\n")

    config = PROJECT / ".vitepress" / "config.mts"
    text = config.read_text(encoding="utf-8")
    if "import enConfig from './config/en'" not in text:
        text = text.replace("import ltConfig from './config/lt'\n", "import ltConfig from './config/lt'\nimport enConfig from './config/en'\n")
    if "link: '/en/'" not in text:
        text = text.replace(
            "    },\n  },",
            "    },\n    en: {\n      label: 'English',\n      lang: 'en',\n      link: '/en/',\n      ...enConfig\n    },\n  },",
        )
    config.write_text(text, encoding="utf-8", newline="\n")


def update_pdf_metadata() -> None:
    path = PROJECT / "pdf-report" / "report-lt.typ"
    text = path.read_text(encoding="utf-8")
    replacements = {
        "2024-2025": "2025-2026",
        "2024–2025": "2025–2026",
        "VU_SA_Ataskaita_2024_2025": "VU_SA_Ataskaita_2025_2026",
        "date: datetime(year: 2025, month: 5, day: 16)": "date: datetime(year: 2026, month: 5, day: 16)",
        'date: "2025 m. gegužė"': 'date: "2026 m. gegužė"',
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    path.write_text(text, encoding="utf-8", newline="\n")

    en_path = PROJECT / "pdf-report" / "report-en.typ"
    en_text = en_path.read_text(encoding="utf-8")
    en_replacements = {
        "2024-2025": "2025-2026",
        "2024–2025": "2025–2026",
        "VU_SR_Report_2024_2025": "VU_SR_Report_2025_2026",
        "date: datetime(year: 2025, month: 5, day: 16)": "date: datetime(year: 2026, month: 5, day: 16)",
        'date: "May 2025"': 'date: "May 2026"',
    }
    for old, new in en_replacements.items():
        en_text = en_text.replace(old, new)
    en_path.write_text(en_text, encoding="utf-8", newline="\n")


def update_pdf_scripts() -> None:
    replacements = {
        "pdf-report/build-en.sh": {
            'VU_SR_Report_2024_2025.pdf': 'VU_SR_Report_2025_2026.pdf',
        },
        "pdf-report/compress-pdf.js": {
            "VU_SR_Report_2024_2025.pdf": "VU_SR_Report_2025_2026.pdf",
        },
        "pdf-report/compress-pdf-advanced.js": {
            "VU_SR_Report_2024_2025.pdf": "VU_SR_Report_2025_2026.pdf",
        },
    }
    for rel, reps in replacements.items():
        path = PROJECT / rel
        text = path.read_text(encoding="utf-8")
        for old, new in reps.items():
            text = text.replace(old, new)
        path.write_text(text, encoding="utf-8", newline="\n")


def main() -> None:
    write_markdown_pages()
    write_json(PROJECT / "src" / "data" / "stats.json", stats_from_doc())
    write_json(PROJECT / "src" / "data" / "timeline.json", timeline_from_doc())
    copy_gallery_and_data()
    export_people_csv()
    update_index()
    write_english_support_pages()
    remove_empty_english_pages()
    update_configs()
    update_pdf_metadata()
    update_pdf_scripts()


if __name__ == "__main__":
    main()
